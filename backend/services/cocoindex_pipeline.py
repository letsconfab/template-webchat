"""CocoIndex pipeline: local cache dir → chunk → embed → entity extract → Neo4j + Qdrant."""

from __future__ import annotations

import json
import logging
import re
from collections.abc import AsyncIterator
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import httpx

import cocoindex as coco
from cocoindex.connectors import localfs, neo4j
from cocoindex.ops.sentence_transformers import SentenceTransformerEmbedder
from cocoindex.ops.text import RecursiveSplitter

from backend.llm_providers import LLMProvider

logger = logging.getLogger(__name__)

# ── Module-level config (set by build_pipeline) ──────────────────────────

_pipeline_cfg: dict[str, Any] = {}

# ── Context keys ──────────────────────────────────────────────────────────

NEO4J_DB = coco.ContextKey[neo4j.ConnectionFactory]("neo4j_db")

# ── Lifespan ──────────────────────────────────────────────────────────────


@coco.lifespan
async def _lifespan(builder: coco.EnvironmentBuilder) -> AsyncIterator[None]:
    conn_factory = neo4j.ConnectionFactory(
        uri=_pipeline_cfg["neo4j_uri"],
        auth=(_pipeline_cfg["neo4j_user"], _pipeline_cfg["neo4j_password"]),
        database=_pipeline_cfg["neo4j_database"],
    )
    builder.provide(NEO4J_DB, conn_factory)
    yield


# ── Schemas for Neo4j ─────────────────────────────────────────────────────

COLLECTION_NAME = "kb_chunks"
VECTOR_SIZE = 384


@dataclass
class Document:
    filename: str
    folder_path: str
    file_type: str
    synced_at: str


@dataclass
class Chunk:
    id: int
    text: str


@dataclass
class Entity:
    name: str
    type: str
    description: str


@dataclass
class Mention:
    snippet: str


@dataclass
class RelatedTo:
    relationship_type: str


# ── Text extraction helpers ───────────────────────────────────────────────

_splitter = RecursiveSplitter()

_SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".html"}


def _extract_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        return "\n\n".join(p.extract_text() for p in reader.pages)
    if ext == ".docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(str(file_path))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return file_path.read_text(encoding="utf-8", errors="ignore")


def _extract_file_type(file_path: Path) -> str:
    ext = file_path.suffix.lower().lstrip(".")
    if ext in ("md", "markdown"):
        return "markdown"
    if ext == "pdf":
        return "pdf"
    if ext == "docx":
        return "docx"
    if ext == "csv":
        return "csv"
    return "txt"


# ── LLM entity extraction ────────────────────────────────────────────────

_ENTITY_EXTRACT_PROMPT = """\
You are extracting structured information from text to build a knowledge graph.

Given the text below, identify:
1. Important entities mentioned (people, organizations, concepts, technologies, locations, dates)
2. The type of each entity (PERSON, ORGANIZATION, CONCEPT, TECHNOLOGY, LOCATION, EVENT)
3. A brief description of each entity
4. Relationships between entities (e.g., "works_for", "part_of", "related_to", "located_in")

Return ONLY valid JSON with this exact structure:
{{
  "entities": [
    {{"name": "Entity Name", "type": "PERSON|ORGANIZATION|CONCEPT|TECHNOLOGY|LOCATION|EVENT", "description": "Brief description"}}
  ],
  "relationships": [
    {{"source": "Entity Name", "target": "Entity Name", "type": "relationship_type"}}
  ]
}}

If no entities are found, return {{"entities": [], "relationships": []}}.

Text:
---
{text}
---
"""


async def _extract_entities_from_text(
    text: str, provider: str, model: str, api_key: str
) -> dict:
    llm = LLMProvider(provider, model, api_key).get_llm()
    if not llm:
        return {"entities": [], "relationships": []}

    from langchain_core.messages import HumanMessage

    truncated = text[:8000]
    prompt = _ENTITY_EXTRACT_PROMPT.format(text=truncated)

    try:
        response = await llm.ainvoke([HumanMessage(content=prompt)])
        match = re.search(r"\{.*\}", str(response.content), re.DOTALL)
        if match:
            return json.loads(match.group())
    except Exception as e:
        logger.warning("Entity extraction LLM call failed: %s", e)

    return {"entities": [], "relationships": []}


# ── Qdrant helpers ────────────────────────────────────────────────────────


async def _ensure_qdrant_collection(qdrant_url: str) -> None:
    async with httpx.AsyncClient() as client:
        resp = await client.get(f"{qdrant_url}/collections/{COLLECTION_NAME}")
        if resp.status_code == 200:
            return
        payload = {
            "name": COLLECTION_NAME,
            "vectors": {"size": VECTOR_SIZE, "distance": "Cosine"},
        }
        resp = await client.put(f"{qdrant_url}/collections/{COLLECTION_NAME}", json=payload)
        if resp.status_code not in (200, 201):
            logger.warning("Qdrant collection creation: %s", resp.text)


async def _push_to_qdrant(
    qdrant_url: str, chunk_id: int, filename: str, text: str, embedding: list[float],
) -> None:
    async with httpx.AsyncClient() as client:
        payload = {
            "points": [{
                "id": chunk_id,
                "vector": embedding,
                "payload": {"filename": filename, "text": text[:5000]},
            }]
        }
        await client.put(f"{qdrant_url}/collections/{COLLECTION_NAME}/points", json=payload)


async def _remove_from_qdrant(qdrant_url: str, filename: str) -> None:
    async with httpx.AsyncClient() as client:
        filter_payload = {
            "filter": {"must": [{"key": "filename", "match": {"value": filename}}]}
        }
        await client.post(
            f"{qdrant_url}/collections/{COLLECTION_NAME}/points/delete", json=filter_payload,
        )


# ── Pipeline functions ────────────────────────────────────────────────────


@coco.fn(memo=True)
async def process_file(
    file: localfs.File,
    doc_table: neo4j.TableTarget[Document],
    chunk_table: neo4j.TableTarget[Chunk],
    entity_table: neo4j.TableTarget[Entity],
    mention_rel: neo4j.RelationTarget[Any],
    related_rel: neo4j.RelationTarget[Any],
) -> None:
    file_path = Path(file.file_path)
    ext = file_path.suffix.lower()

    logger.info("process_file called: %s (ext=%s)", file_path, ext)

    if ext not in _SUPPORTED_EXTENSIONS:
        logger.info("  -> unsupported extension, skipping")
        return

    text = _extract_text(file_path)
    if not text.strip():
        logger.info("  -> empty text, skipping")
        return

    file_type = _extract_file_type(file_path)
    filename = file_path.name
    folder = str(file_path.parent)
    synced_at = datetime.utcnow().isoformat()

    doc_table.declare_record(
        row=Document(
            filename=filename, folder_path=folder, file_type=file_type, synced_at=synced_at,
        )
    )

    chunks = _splitter.split(text, chunk_size=2000, chunk_overlap=500, language="markdown")

    # Read config from module-level dict
    embedder: SentenceTransformerEmbedder = _pipeline_cfg["embedder"]
    qdrant_url: str = _pipeline_cfg["qdrant_url"]
    llm_provider: str = _pipeline_cfg["llm_provider"]
    llm_model: str = _pipeline_cfg["llm_model"]
    llm_api_key: str = _pipeline_cfg["llm_api_key"]

    await _ensure_qdrant_collection(qdrant_url)
    await _remove_from_qdrant(qdrant_url, filename)

    for idx, chunk in enumerate(chunks):
        chunk_id = hash(f"{filename}:{idx}") % (2**31)

        embedding = await embedder.embed(chunk.text)
        embedding_list = embedding.tolist() if hasattr(embedding, "tolist") else list(embedding)

        await _push_to_qdrant(qdrant_url, chunk_id, filename, chunk.text, embedding_list)

        chunk_table.declare_record(row=Chunk(id=chunk_id, text=chunk.text))

    if llm_api_key:
        extracted = await _extract_entities_from_text(
            text, llm_provider, llm_model, llm_api_key
        )
        for ent in extracted.get("entities", []):
            entity_table.declare_record(
                row=Entity(
                    name=ent["name"],
                    type=ent.get("type", "CONCEPT"),
                    description=ent.get("description", ""),
                )
            )
        for rel in extracted.get("relationships", []):
            related_rel.declare_relation(
                from_id=rel["source"],
                to_id=rel["target"],
                record=RelatedTo(relationship_type=rel["type"]),
            )


# ── App builder ───────────────────────────────────────────────────────────


def build_pipeline(
    cache_dir: str,
    neo4j_uri: str,
    neo4j_user: str,
    neo4j_password: str,
    neo4j_database: str,
    qdrant_url: str,
    embedding_model: str,
    llm_provider: str,
    llm_model: str,
    llm_api_key: str,
) -> coco.App:
    """Build the CocoIndex pipeline application."""
    _pipeline_cfg.update(
        neo4j_uri=neo4j_uri,
        neo4j_user=neo4j_user,
        neo4j_password=neo4j_password,
        neo4j_database=neo4j_database,
        qdrant_url=qdrant_url,
        llm_provider=llm_provider,
        llm_model=llm_model,
        llm_api_key=llm_api_key,
    )
    _pipeline_cfg.pop("embedder", None)

    @coco.fn
    async def app_main() -> None:
        doc_table = await neo4j.mount_table_target(
            NEO4J_DB,
            "Document",
            await neo4j.TableSchema.from_class(Document, primary_key="filename"),
            primary_key="filename",
        )
        chunk_table = await neo4j.mount_table_target(
            NEO4J_DB,
            "Chunk",
            await neo4j.TableSchema.from_class(Chunk, primary_key="id"),
            primary_key="id",
        )
        entity_table = await neo4j.mount_table_target(
            NEO4J_DB,
            "Entity",
            await neo4j.TableSchema.from_class(Entity, primary_key="name"),
            primary_key="name",
        )
        mention_rel = await neo4j.mount_relation_target(
            NEO4J_DB, "MENTIONS", chunk_table, entity_table,
        )
        related_rel = await neo4j.mount_relation_target(
            NEO4J_DB, "RELATED_TO", entity_table, entity_table,
        )

        embedder = SentenceTransformerEmbedder(embedding_model)
        _pipeline_cfg["embedder"] = embedder

        source = localfs.walk_dir(cache_dir, recursive=True)
        await coco.mount_each(
            process_file, source.items(),
            doc_table, chunk_table, entity_table, mention_rel, related_rel,
        )

    return coco.App(
        coco.AppConfig(name="WebChatKBIndexer"),
        app_main,
    )
